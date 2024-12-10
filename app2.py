from flask import Flask, render_template, request, make_response, send_file
from transformers import pipeline
import httpx
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor
from flask_caching import Cache
from io import BytesIO
from docx import Document
import threading
import os

app = Flask(__name__)

# Enable caching
cache = Cache(app, config={"CACHE_TYPE": "simple"})

# Initialize the summarizer pipeline with GPU if available
device = 0 if os.getenv("CUDA_VISIBLE_DEVICES") else -1
summarizer = pipeline("summarization", model="facebook/bart-large-cnn", device=device)

# Create an HTTP client for persistent connections
http_client = httpx.Client(timeout=10.0, headers={"Connection": "keep-alive"})

# Function to split text into manageable chunks
def split_text(text, max_words=500):
    words = text.split()
    for i in range(0, len(words), max_words):
        yield ' '.join(words[i:i + max_words])

# Optimized function for summarizing text chunks in parallel
def summarize_chunk(chunk):
    return summarizer(chunk, max_length=150, min_length=50, do_sample=False)[0]["summary_text"]

@cache.memoize(timeout=3600)
def summarize_text(url):
    try:
        # Fetch webpage content (cached to avoid redundant downloads)
        response = http_client.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')

        # Extract text from <p> tags
        text = ' '.join(p.get_text() for p in soup.find_all('p'))

        # Summarize the text in chunks using a thread pool
        with ThreadPoolExecutor(max_workers=os.cpu_count()) as executor:
            summaries = list(executor.map(summarize_chunk, split_text(text)))

        return ' '.join(summaries)
    except Exception as e:
        return f"An error occurred: {e}"

# Word Document Generation Function (thread-safe)
def generate_word(summary):
    doc = Document()
    doc.add_heading('Summary', 0)
    doc.add_paragraph(summary)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Route for the homepage
@app.route("/", methods=["GET", "POST"])
def index():
    summary = ""
    if request.method == "POST":
        url = request.form["url"]
        # Summarize the content
        summary = summarize_text(url)

        # Generate Word document asynchronously
        threading.Thread(target=generate_word, args=(summary,)).start()

    response = make_response(render_template("index.html", summary=summary))
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

# Route to download Word
@app.route("/download_word", methods=["GET"])
def download_word():
    summary = request.args.get('summary')
    word_buffer = generate_word(summary)
    return send_file(word_buffer, as_attachment=True, download_name="summary.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == "__main__":
    app.run(debug=True)
